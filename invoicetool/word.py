# import subprocess
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterator

from docx import Document

TEMPORARY_MS_WORD_DOCUMENT_BYTES = b"\x15Microsoft Office User\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x15\x00M\x00i\x00c\x00r\x00o\x00s\x00o\x00f\x00t\x00 \x00O\x00f\x00f\x00i\x00c\x00e\x00 \x00U\x00s\x00e\x00r\x00\x00\x00k\x00t\x00o\x00p\x00/\x00A\x00c\x00c\x00o\x00u\x00n\x00t\x00s\x00 \x002\x000\x002\x000\x00/\x00F\x00e\x00b\x00 \x00I\x00n\x00v\x00o\x00i\x00c\x00e\x00s\x00"


@dataclass
class WordDocument:
    name: str
    modification_time: float | int
    size: int
    hash: str
    original_path: str
    text: str
    _hash_type: str
    _path: str

    def __str__(self):
        return f"WordDocument(name={self.name}, size={self.size} B, modification_time={self.modification_time})"

    def to_dict(self):
        return asdict(self)

    @classmethod
    def from_dict(cls, d):
        return cls(**d)


def get_paragraphs(doc: Document) -> Iterator[str]:
    """get all paragraphs from a docx document"""
    for paragraph in doc.paragraphs:
        p = paragraph.text.strip()
        if not p:
            continue
        yield " ".join(p.split())


# def extract_text_from_document_as_list(filepath: Path) -> list[str]:
#     """Extract all text from a document and return the text as a list of paragraphs"""
#     if filepath.suffix == ".docx":
#         return extract_text_from_docx_as_list(filepath)
#     elif filepath.suffix == ".doc":
#         return extract_text_from_doc_as_list(filepath)
#     else:
#         raise ValueError(f"Unsupported file extension: {filepath.suffix}")


# def extract_text_from_docx_as_list(filepath: Path) -> list[str]:
#     """Extract all text from a `.docx` file and return a list of paragraphs"""
#     text = get_text_from_docx(filepath)
#     return text_to_paragraphs(text)


# def extract_text_from_doc_as_list(filepath: Path) -> list[str]:
#     """Extract all text from a `.doc` file and return the text"""
#     text = extract_text_from_doc(filepath)
#     return text_to_paragraphs(text)


# def text_to_paragraphs(text: str) -> list[str]:
#     """Convert a string of text into a list of paragraphs"""
#     return [p for p in text.split("\n") if p]


# def extract_text_from_document(filepath: Path) -> str:
#     """Extract all text from a document and return the text"""
#     if filepath.suffix == ".docx":
#         return get_text_from_docx(filepath)
#     elif filepath.suffix == ".doc":
#         return extract_text_from_doc(filepath)
#     else:
#         raise ValueError(f"Unsupported file extension: {filepath.suffix}")


def get_text_from_docx(filepath: Path | str) -> str:
    """Extract all text from a `.docx` file and return the text"""
    doc = Document(filepath)
    text = "\n".join(get_paragraphs(doc))
    return text


# def extract_text_from_doc(filepath: Path) -> str:
#     """Extract all text from a `.doc` file and return the text"""
#     # process is the completed process
#     process = subprocess.run(
#         ["antiword", filepath.as_posix()],
#         capture_output=True,
#         text=True,
#     )
#     raw_text = process.stdout
#     return raw_text.strip()
