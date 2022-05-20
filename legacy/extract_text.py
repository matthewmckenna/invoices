#!/usr/bin/env python
"""dump all text from a word document to a database"""
import argparse
from pathlib import Path
import subprocess
from typing import Iterable

from docx import Document

from get_duplicates import scantree


def get_paragraphs(doc: Document) -> Iterable[str]:
    """get all paragraphs from a document"""
    for paragraph in doc.paragraphs:
        p = paragraph.text.strip()

        # get rid of blank lines
        if not p:
            continue

        # TODO: maybe call generic `sanitise` which
        # does multiple cleaning operations
        yield multi_whitespace_to_space(p)


def multi_whitespace_to_space(s: str) -> str:
    """squeeze multiple whitespace characters into a single space"""
    return ' '.join(s.split())


def get_all_text(filepath: Path) -> str:
    """accepts a filepath and returns all text from the docx file"""
    doc = Document(filepath)
    text = '\n'.join(get_paragraphs(doc))
    return text


def get_all_text_list(filepath: Path) -> list[str]:
    """accepts a filepath and returns all text from the docx file
    as a list of paragraphs"""
    doc = Document(filepath)
    text_list = list(get_paragraphs(doc))
    return text_list


def main(args):
    """main entry point for the script"""
    directory = Path(args.directory)

    # TODO: configure database

    # TODO: get this from CLI argument
    # extensions that we're interested in
    # extensions = {'.doc', '.docx'}
    extensions = {'.docx'}

    for entry in scantree(directory):
        if entry.is_file():
            # create a path object
            p = Path(entry.path)

            # skip files we're not interested in
            if p.suffix not in extensions:
                continue

            text = extract_docx(p)

            # TODO: add text to database here


def extract_doc(filepath: Path):
    # process is the completed process
    process = subprocess.run(
        ['antiword', filepath.as_posix()],
        capture_output=True,
        text=True,
    )
    # TODO: process the raw text and strip whitespace, newlines, etc.
    raw_text = process.stdout
    return raw_text


def extract_docx(filepath: Path) -> str:
    """extract all text from a .docx file and return the text"""
    doc = Document(filepath)
    text = '\n'.join(get_paragraphs(doc))
    return text


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='extract text from all Word files in a given directory',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument('directory', help='starting directory')
    args = parser.parse_args()
    main(args)
