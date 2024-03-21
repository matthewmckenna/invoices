import io
from pathlib import Path

import pytest
from docx import Document

from invoicetool.config import Config


@pytest.fixture
def config(tmp_path_factory) -> Config:
    """Return a default config object"""
    output_dir = tmp_path_factory.mktemp("output_dir")
    return Config(
        extensions={".doc", ".docx"},
        output_directory=output_dir,
        hash_function_algorithm="sha1",
    )


@pytest.fixture(scope="session")
def invoices_dir(tmp_path_factory):
    """Create a simple directory structure with different filetypes.

    Directory structure:
    .
    ├── another-level
    │   ├── document03.doc
    │   └── spreadsheet02.xlsx
    ├── document01.doc
    ├── document02.docx
    └── spreadsheet01.xls

    """
    invoices_dir = tmp_path_factory.mktemp("invoices")
    another_level = invoices_dir / "another-level"
    another_level.mkdir()

    for filepath in [
        invoices_dir / "document01.doc",
        invoices_dir / "document02.docx",
        invoices_dir / "spreadsheet01.xls",
        another_level / "document03.doc",
        another_level / "spreadsheet02.xlsx",
    ]:
        filepath.touch()

    return invoices_dir


@pytest.fixture(scope="function")
def empty_directory(tmp_path_factory: pytest.TempPathFactory) -> Path:
    empty_dir = tmp_path_factory.mktemp("empty_directory")
    return empty_dir


@pytest.fixture(scope="function")
def non_empty_directory(tmp_path: Path) -> Path:
    (tmp_path / "file.txt").touch()
    return tmp_path


@pytest.fixture(scope="function")
def nested_empty_directories(tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("nested")
    (path / "empty1" / "empty2").mkdir(parents=True, exist_ok=True)
    return path


@pytest.fixture
def docx_bytes() -> io.BytesIO:
    """Create a new in-memory document"""
    with io.BytesIO() as document_buffer:
        doc = Document()
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph.")
        doc.save(document_buffer)
        document_bytes = document_buffer.getvalue()

    return io.BytesIO(document_bytes)


@pytest.fixture
def documents_directory():
    project_directory = Config.project_directory
    data_directory = project_directory / "data"
    return data_directory / "sample-documents"


@pytest.fixture
def real_doc_filepath(documents_directory: Path) -> Path:
    return documents_directory / "older-format.doc"


@pytest.fixture
def real_docx_filepath(documents_directory: Path) -> Path:
    return documents_directory / "newer-format.docx"
